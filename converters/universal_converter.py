import os
import asyncio
import subprocess
import tempfile
from pathlib import Path
import logging
import html
from config import Config

logger = logging.getLogger(__name__)

class UniversalConverter:
    def __init__(self):
        self.supported_formats = {}
        for category, formats in Config.SUPPORTED_FORMATS.items():
            for fmt in formats:
                self.supported_formats[fmt] = category
    
    async def _check_ffmpeg_available(self):
        """Check if FFmpeg is available"""
        try:
            process = await asyncio.create_subprocess_exec(
                'ffmpeg', '-version',
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE
            )
            stdout, stderr = await process.communicate()
            return process.returncode == 0
        except:
            return False

    async def convert_file(self, input_path, output_format, input_extension=None):
        """Professional file conversion with high quality"""
        try:
            if not input_extension:
                input_extension = os.path.splitext(input_path)[1].lstrip('.').lower()
            
            logger.info(f"Converting {input_extension} to {output_format}")
            
            # Get file categories
            input_category = self.supported_formats.get(input_extension)
            output_category = self.supported_formats.get(output_format)
            
            if not input_category:
                raise Exception(f"Unsupported input format: {input_extension}")
            if not output_category:
                raise Exception(f"Unsupported output format: {output_format}")
            
            # Route to appropriate converter
            if input_category == 'image':
                return await self._convert_image(input_path, output_format, input_extension)
            elif input_category == 'audio':
                return await self._convert_audio(input_path, output_format, input_extension)
            elif input_category == 'video':
                return await self._convert_video(input_path, output_format, input_extension)
            elif input_category == 'document':
                return await self._convert_document(input_path, output_format, input_extension)
            elif input_category == 'presentation':
                return await self._convert_presentation(input_path, output_format, input_extension)
            else:
                raise Exception(f"No converter for category: {input_category}")
                
        except Exception as e:
            logger.error(f"Universal conversion error: {e}")
            raise
    
    async def _convert_image(self, input_path, output_format, input_extension):
        """Professional image conversion with high quality"""
        try:
            from PIL import Image, ImageSequence
            
            # Create proper output path with correct extension
            base_name = os.path.splitext(input_path)[0]
            output_path = f"{base_name}_converted.{output_format}"
            
            logger.info(f"Image conversion: {input_extension} -> {output_format}")
            logger.info(f"Output path: {output_path}")
            
            # Handle GIF conversions specially
            if input_extension.lower() == 'gif':
                return await self._convert_gif_to_static(input_path, output_path, output_format)
            
            with Image.open(input_path) as img:
                # Handle format-specific conversions with professional settings
                if output_format in ['jpg', 'jpeg']:
                    # Professional JPEG conversion
                    if img.mode in ('RGBA', 'LA', 'P'):
                        if img.mode == 'P' and 'transparency' in img.info:
                            img = img.convert('RGBA')
                        background = Image.new('RGB', img.size, (255, 255, 255))
                        if img.mode == 'RGBA':
                            background.paste(img, mask=img.split()[-1])
                        else:
                            background.paste(img)
                        img = background
                    elif img.mode != 'RGB':
                        img = img.convert('RGB')
                    
                    # High quality JPEG save
                    img.save(output_path, 'JPEG', quality=95, optimize=True, progressive=True)
                    
                elif output_format == 'png':
                    # High quality PNG with optimization
                    if img.mode == 'P':
                        img = img.convert('RGBA')
                    elif img.mode == 'RGB':
                        # Ensure PNG maintains transparency capability
                        img = img.convert('RGBA')
                    img.save(output_path, 'PNG', optimize=True)
                    
                elif output_format == 'bmp':
                    # BMP conversion
                    if img.mode != 'RGB':
                        img = img.convert('RGB')
                    img.save(output_path, 'BMP')
                    
                elif output_format == 'pdf':
                    # Professional image to PDF conversion
                    return await self._image_to_pdf_advanced(input_path, output_path)
                elif output_format == 'gif':
                    # Convert to animated GIF
                    return await self._convert_to_animated_gif(input_path, output_path, input_extension)
                else:
                    # Fallback for other formats
                    img.save(output_path, format=output_format.upper())
                
                # Verify the output file was created with correct format
                if os.path.exists(output_path):
                    logger.info(f"Image conversion successful: {output_path}")
                    return output_path
                else:
                    raise Exception("Output file was not created")
                
        except Exception as e:
            logger.error(f"Image conversion error: {e}")
            raise Exception(f"Professional image conversion failed: {str(e)}")
    
    async def _convert_gif_advanced(self, input_path, output_path, input_ext, output_format):
        """Advanced GIF conversion with optimized quality and performance"""
        try:
            from PIL import Image, ImageSequence, ImageEnhance
            
            # Handle GIF to other formats conversion
            if input_ext.lower() == 'gif' and output_format != 'gif':
                return await self._convert_gif_to_static(input_path, output_path, output_format)
            
            # Handle other formats to GIF conversion
            elif output_format == 'gif':
                return await self._convert_to_animated_gif(input_path, output_path, input_ext)
            
            return output_path
            
        except Exception as e:
            logger.error(f"GIF conversion error: {e}")
            raise Exception(f"GIF conversion failed: {str(e)}")

    async def _convert_gif_to_static(self, input_path, output_path, output_format):
        """Convert animated GIF to static image formats"""
        try:
            from PIL import Image, ImageSequence
            
            with Image.open(input_path) as img:
                # Extract first frame for static conversion
                first_frame = None
                for frame in ImageSequence.Iterator(img):
                    first_frame = frame.copy()
                    break
                
                if not first_frame:
                    raise Exception("No frames found in GIF")
                
                # Handle format-specific conversions
                if output_format in ['jpg', 'jpeg']:
                    if first_frame.mode != 'RGB':
                        first_frame = first_frame.convert('RGB')
                    first_frame.save(output_path, 'JPEG', quality=95, optimize=True)
                    
                elif output_format == 'png':
                    if first_frame.mode == 'P':
                        first_frame = first_frame.convert('RGBA')
                    elif first_frame.mode == 'RGB':
                        first_frame = first_frame.convert('RGBA')
                    first_frame.save(output_path, 'PNG', optimize=True)
                    
                elif output_format == 'bmp':
                    if first_frame.mode != 'RGB':
                        first_frame = first_frame.convert('RGB')
                    first_frame.save(output_path, 'BMP')
                    
                elif output_format == 'pdf':
                    return await self._image_to_pdf_advanced(input_path, output_path)
                else:
                    first_frame.save(output_path, format=output_format.upper())
                
                return output_path
                
        except Exception as e:
            raise Exception(f"GIF to {output_format} conversion failed: {str(e)}")

    async def _convert_to_animated_gif(self, input_path, output_path, input_ext):
        """Convert various formats to animated GIF"""
        try:
            # For video formats, use FFmpeg for better quality
            if input_ext in ['mp4', 'avi', 'mov', 'mkv']:
                return await self._convert_video_to_gif_ffmpeg(input_path, output_path)
            
            # For image formats, create animated GIF
            elif input_ext in ['jpg', 'jpeg', 'png', 'bmp']:
                return await self._convert_image_to_animated_gif(input_path, output_path)
            
            else:
                raise Exception(f"Conversion from {input_ext} to GIF not supported")
                
        except Exception as e:
            raise Exception(f"Animated GIF creation failed: {str(e)}")

    async def _convert_image_to_animated_gif(self, input_path, output_path):
        """Convert static image to animated GIF with effects"""
        try:
            from PIL import Image, ImageEnhance, ImageFilter
            
            with Image.open(input_path) as img:
                # Convert to RGB if necessary
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                
                frames = []
                total_frames = 10  # Reduced frames for better performance
                base_duration = 150  # Base duration in ms
                
                # Create different variations for animation
                for i in range(total_frames):
                    frame = img.copy()
                    
                    # Add subtle animation effects
                    if i % 5 == 0:  # Every 5th frame
                        # Slight brightness variation
                        enhancer = ImageEnhance.Brightness(frame)
                        frame = enhancer.enhance(1.05)
                    elif i % 3 == 0:  # Every 3rd frame
                        # Slight contrast variation
                        enhancer = ImageEnhance.Contrast(frame)
                        frame = enhancer.enhance(1.02)
                    
                    frames.append(frame)
                
                # Calculate optimized duration
                total_duration = 2000  # 2 seconds total
                frame_duration = total_duration // len(frames)
                
                # Save as optimized animated GIF
                frames[0].save(
                    output_path,
                    format='GIF',
                    save_all=True,
                    append_images=frames[1:],
                    duration=frame_duration,
                    loop=0,
                    optimize=True,
                    disposal=2  # Background disposal
                )
                
                # Optimize file size
                await self._optimize_gif_size(output_path)
                
                return output_path
                
        except Exception as e:
            raise Exception(f"Image to animated GIF conversion failed: {str(e)}")

    async def _convert_video_to_gif_ffmpeg(self, input_path, output_path):
        """Convert video to high-quality optimized GIF using FFmpeg"""
        try:
            # First pass: analyze video duration
            probe_cmd = [
                'ffprobe', '-v', 'error', '-show_entries', 'format=duration',
                '-of', 'default=noprint_wrappers=1:nokey=1', input_path
            ]
            
            process = await asyncio.create_subprocess_exec(
                *probe_cmd,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE
            )
            stdout, stderr = await process.communicate()
            
            duration = float(stdout.decode().strip()) if stdout else 3.0
            # Limit to 5 seconds maximum for GIF
            duration = min(duration, 5.0)
            
            # Use optimized FFmpeg settings for GIF creation
            cmd = [
                'ffmpeg', '-i', input_path,
                '-t', str(duration),  # Use actual duration or max 5 seconds
                '-vf', 'fps=10,scale=500:-1:flags=lanczos,split[s0][s1];[s0]palettegen=stats_mode=diff[p];[s1][p]paletteuse=dither=bayer:bayer_scale=3',
                '-y', output_path
            ]
            
            await self._run_command(cmd, timeout=120)
            
            # Post-process optimization
            await self._optimize_gif_size(output_path)
            
            return output_path
            
        except Exception as e:
            logger.warning(f"FFmpeg GIF conversion failed, trying fallback: {e}")
            # Fallback to simpler conversion
            return await self._convert_video_to_gif_fallback(input_path, output_path)

    async def _convert_video_to_gif_fallback(self, input_path, output_path):
        """Fallback video to GIF conversion"""
        try:
            cmd = [
                'ffmpeg', '-i', input_path,
                '-t', '3',  # 3 seconds
                '-vf', 'fps=8,scale=400:-1:flags=lanczos',
                '-y', output_path
            ]
            
            await self._run_command(cmd, timeout=90)
            return output_path
            
        except Exception as e:
            raise Exception(f"Fallback GIF conversion failed: {str(e)}")

    async def _optimize_gif_size(self, gif_path):
        """Optimize GIF file size using various techniques"""
        try:
            from PIL import Image, ImageSequence
            
            # Check if file size is reasonable (under 8MB)
            file_size = os.path.getsize(gif_path)
            if file_size <= 8 * 1024 * 1024:  # 8MB
                return gif_path
            
            logger.info(f"Optimizing GIF size: {file_size / (1024 * 1024):.1f}MB")
            
            # Re-encode with PIL for better compression
            with Image.open(gif_path) as img:
                frames = []
                durations = []
                
                # Extract frames and durations
                for frame in ImageSequence.Iterator(img):
                    frames.append(frame.copy())
                    # Get frame duration (default to 100ms if not available)
                    try:
                        duration = frame.info.get('duration', 100)
                        durations.append(duration)
                    except:
                        durations.append(100)
                
                # Reduce number of frames if too many
                if len(frames) > 30:
                    skip_factor = len(frames) // 15  # Target 15 frames
                    frames = frames[::skip_factor]
                    durations = durations[::skip_factor]
                    frames = frames[:15]  # Max 15 frames
                    durations = durations[:15]
                
                # Save with optimized settings
                frames[0].save(
                    gif_path,
                    format='GIF',
                    save_all=True,
                    append_images=frames[1:],
                    duration=durations[0] if durations else 100,
                    loop=0,
                    optimize=True,
                    disposal=2
                )
            
            optimized_size = os.path.getsize(gif_path)
            logger.info(f"GIF optimized: {optimized_size / (1024 * 1024):.1f}MB")
            
            return gif_path
            
        except Exception as e:
            logger.warning(f"GIF optimization failed: {e}")
            return gif_path  # Return original if optimization fails
    
    async def _image_to_pdf_advanced(self, input_path, output_path):
        """Professional image to PDF conversion"""
        try:
            from PIL import Image
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.utils import ImageReader
            
            with Image.open(input_path) as img:
                # Use ReportLab for professional PDF creation
                c = canvas.Canvas(output_path, pagesize=letter)
                width, height = letter
                
                # Calculate scaling to fit image on page
                img_width, img_height = img.size
                scale_x = width / img_width
                scale_y = height / img_height
                scale = min(scale_x, scale_y) * 0.9  # 90% of page with margin
                
                # Calculate position to center image
                x = (width - (img_width * scale)) / 2
                y = (height - (img_height * scale)) / 2
                
                # Draw image
                img_temp_path = input_path + '_temp.jpg'
                img.convert('RGB').save(img_temp_path, 'JPEG', quality=90)
                c.drawImage(img_temp_path, x, y, img_width * scale, img_height * scale)
                c.save()
                
                # Cleanup temp file
                if os.path.exists(img_temp_path):
                    os.remove(img_temp_path)
                    
            return output_path
        except Exception as e:
            raise Exception(f"Professional image to PDF conversion failed: {str(e)}")
    
    async def _convert_audio(self, input_path, output_format, input_extension):
        """Convert audio files using FFmpeg with smart compression"""
        try:
            # Check FFmpeg availability
            if not await self._check_ffmpeg_available():
                raise Exception("FFmpeg is required for audio conversion but is not installed")
            
            output_path = input_path.rsplit('.', 1)[0] + f'.{output_format}'
            
            # Get input file size to determine optimal settings
            input_size = os.path.getsize(input_path)
            input_size_mb = input_size / (1024 * 1024)
            
            logger.info(f"Audio conversion: {input_extension} -> {output_format}, Input size: {input_size_mb:.1f}MB")
            
            # Smart compression based on input size
            if input_size_mb > 50:  # Large files
                compression_level = 'high'
            elif input_size_mb > 20:  # Medium files
                compression_level = 'medium'
            else:  # Small files
                compression_level = 'low'
            
            cmd = [
                'ffmpeg', '-i', input_path,
                '-y',  # Overwrite
                '-loglevel', 'error',
                '-hide_banner',
            ]
            
            # Smart audio codec settings based on format and file size
            if output_format == 'mp3':
                if compression_level == 'high':
                    cmd.extend(['-codec:a', 'libmp3lame', '-b:a', '128k', '-compression_level', '0'])
                elif compression_level == 'medium':
                    cmd.extend(['-codec:a', 'libmp3lame', '-b:a', '192k', '-compression_level', '0'])
                else:
                    cmd.extend(['-codec:a', 'libmp3lame', '-b:a', '256k', '-q:a', '0'])
                    
            elif output_format == 'wav':
                # For WAV, we can control size by using different sample rates and bit depths
                if compression_level == 'high':
                    cmd.extend(['-codec:a', 'pcm_s16le', '-ac', '1', '-ar', '22050'])  # Mono, lower sample rate
                elif compression_level == 'medium':
                    cmd.extend(['-codec:a', 'pcm_s16le', '-ac', '1', '-ar', '44100'])  # Mono, standard sample rate
                else:
                    cmd.extend(['-codec:a', 'pcm_s16le', '-ac', '2', '-ar', '44100'])  # Stereo, standard
                    
            elif output_format == 'aac':
                if compression_level == 'high':
                    cmd.extend(['-codec:a', 'aac', '-b:a', '128k', '-ac', '1'])
                elif compression_level == 'medium':
                    cmd.extend(['-codec:a', 'aac', '-b:a', '192k', '-ac', '2'])
                else:
                    cmd.extend(['-codec:a', 'aac', '-b:a', '256k', '-ac', '2'])
            
            cmd.append(output_path)
            
            logger.info(f"Audio conversion command: {' '.join(cmd)}")
            
            process = await asyncio.create_subprocess_exec(
                *cmd,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE
            )
            
            # Monitor progress with timeout
            stdout, stderr = await asyncio.wait_for(process.communicate(), timeout=300)  # 5 minutes timeout
            
            if process.returncode == 0 and os.path.exists(output_path):
                # Verify output file is valid and check size
                output_size = os.path.getsize(output_path)
                output_size_mb = output_size / (1024 * 1024)
                
                logger.info(f"Audio conversion successful. Output size: {output_size_mb:.1f}MB")
                
                if output_size == 0:
                    raise Exception("Conversion produced empty file")
                
                # If output is still too large, apply additional compression
                if output_size > 45 * 1024 * 1024:  # Over 45MB
                    logger.info(f"Output file too large ({output_size_mb:.1f}MB), applying additional compression")
                    compressed_path = await self._compress_audio_file(output_path, output_format)
                    if compressed_path:
                        return compressed_path
                
                return output_path
            else:
                error_msg = stderr.decode('utf-8', errors='ignore') if stderr else "Audio conversion failed"
                logger.error(f"Audio conversion error: {error_msg}")
                raise Exception(f"Audio conversion error: {error_msg}")
                
        except asyncio.TimeoutError:
            raise Exception("Audio conversion timeout - file might be too large")
        except Exception as e:
            logger.error(f"Audio conversion failed: {str(e)}")
            raise Exception(f"Audio conversion failed: {str(e)}")

    async def _compress_audio_file(self, input_path, output_format):
        """Apply additional compression to audio files that are too large"""
        try:
            compressed_path = input_path.replace(f'.{output_format}', f'_compressed.{output_format}')
            
            logger.info(f"Applying additional compression to reduce file size")
            
            cmd = ['ffmpeg', '-i', input_path, '-y']
            
            if output_format == 'wav':
                # Maximum compression for WAV: mono, low sample rate
                cmd.extend(['-codec:a', 'pcm_s16le', '-ac', '1', '-ar', '16000'])
            elif output_format == 'mp3':
                # Maximum compression for MP3: low bitrate
                cmd.extend(['-codec:a', 'libmp3lame', '-b:a', '96k'])
            elif output_format == 'aac':
                # Maximum compression for AAC: low bitrate, mono
                cmd.extend(['-codec:a', 'aac', '-b:a', '96k', '-ac', '1'])
            else:
                cmd.extend(['-codec:a', 'copy'])  # Just copy if we can't compress further
            
            cmd.append(compressed_path)
            
            process = await asyncio.create_subprocess_exec(
                *cmd,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE
            )
            
            stdout, stderr = await asyncio.wait_for(process.communicate(), timeout=120)
            
            if process.returncode == 0 and os.path.exists(compressed_path):
                compressed_size = os.path.getsize(compressed_path)
                compressed_size_mb = compressed_size / (1024 * 1024)
                logger.info(f"Compression successful. New size: {compressed_size_mb:.1f}MB")
                
                # Replace original with compressed version
                os.remove(input_path)
                return compressed_path
            else:
                logger.warning("Additional compression failed, returning original file")
                return input_path
                
        except Exception as e:
            logger.error(f"Audio compression failed: {e}")
            return input_path  # Return original if compression fails
    
    async def _convert_video(self, input_path, output_format, input_extension):
        """Professional video conversion"""
        try:
            if not await self._check_ffmpeg_available():
                raise Exception("FFmpeg is required for video conversion but is not installed")
            
            output_path = input_path.rsplit('.', 1)[0] + f'.{output_format}'
            
            cmd = [
                'ffmpeg', '-i', input_path,
                '-y',
                '-loglevel', 'error',
                '-hide_banner',
            ]
            
            # Professional video conversion settings
            if output_format == 'mp4':
                cmd.extend([
                    '-c:v', 'libx264', '-preset', 'medium', '-crf', '23',
                    '-c:a', 'aac', '-b:a', '192k',
                    '-movflags', '+faststart'
                ])
            elif output_format == 'avi':
                cmd.extend([
                    '-c:v', 'mpeg4', '-qscale:v', '3',
                    '-c:a', 'mp3', '-b:a', '192k'
                ])
            elif output_format == 'mov':
                cmd.extend([
                    '-c:v', 'libx264', '-preset', 'medium', '-crf', '23',
                    '-c:a', 'aac', '-b:a', '192k'
                ])
            elif output_format == 'mkv':
                cmd.extend([
                    '-c:v', 'libx264', '-preset', 'medium', '-crf', '23',
                    '-c:a', 'aac', '-b:a', '192k'
                ])
            elif output_format == 'gif':
                # Use our enhanced GIF conversion for videos
                return await self._convert_video_to_gif_ffmpeg(input_path, output_path)
            
            if output_format != 'gif':  # Already handled for GIF
                cmd.append(output_path)
            
            process = await asyncio.create_subprocess_exec(
                *cmd,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE
            )
            
            stdout, stderr = await asyncio.wait_for(process.communicate(), timeout=300)
            
            if process.returncode == 0 and os.path.exists(output_path):
                return output_path
            else:
                error_msg = stderr.decode('utf-8', errors='ignore') if stderr else "Video conversion failed"
                raise Exception(f"Professional video conversion error: {error_msg}")
                
        except Exception as e:
            raise Exception(f"Professional video conversion failed: {str(e)}")
    
    async def _convert_document(self, input_path, output_format, input_extension):
        """Professional document conversion with high accuracy"""
        try:
            output_path = input_path.rsplit('.', 1)[0] + f'.{output_format}'
            
            logger.info(f"Converting document: {input_extension} -> {output_format}")
            
            # PDF conversions
            if input_extension == 'pdf':
                if output_format == 'txt':
                    return await self._pdf_to_text_advanced(input_path, output_path)
                elif output_format in ['jpg', 'jpeg', 'png']:
                    return await self._pdf_to_images_advanced(input_path, output_path, output_format)
                elif output_format == 'docx':
                    return await self._pdf_to_docx_advanced(input_path, output_path)
                elif output_format == 'xlsx':
                    return await self._pdf_to_excel_advanced(input_path, output_path)
            
            # Text conversions
            elif input_extension == 'txt':
                if output_format == 'pdf':
                    return await self._text_to_pdf_advanced(input_path, output_path)
                elif output_format == 'docx':
                    return await self._text_to_docx_advanced(input_path, output_path)
                elif output_format == 'xlsx':
                    return await self._text_to_excel_advanced(input_path, output_path)
            
            # Word document conversions
            elif input_extension == 'docx':
                if output_format == 'pdf':
                    return await self._docx_to_pdf_advanced(input_path, output_path)
                elif output_format == 'txt':
                    return await self._docx_to_text_advanced(input_path, output_path)
                elif output_format == 'xlsx':
                    return await self._docx_to_excel_advanced(input_path, output_path)
            
            # Excel conversions
            elif input_extension == 'xlsx':
                if output_format == 'pdf':
                    return await self._excel_to_pdf_advanced(input_path, output_path)
                elif output_format == 'txt':
                    return await self._excel_to_text_advanced(input_path, output_path)
                elif output_format == 'docx':
                    return await self._excel_to_docx_advanced(input_path, output_path)
            
            # ODT conversions
            elif input_extension == 'odt':
                if output_format == 'pdf':
                    return await self._odt_to_pdf_advanced(input_path, output_path)
            
            raise Exception(f"Document conversion from {input_extension} to {output_format} not implemented")
            
        except Exception as e:
            logger.error(f"Document conversion error: {e}")
            raise Exception(f"Professional document conversion failed: {str(e)}")
    
    async def _pdf_to_text_advanced(self, input_path, output_path):
        """Advanced PDF to text conversion with formatting preservation"""
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(input_path)
            text_content = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text("text")  # Use "text" for better formatting
                if text.strip():
                    text_content += f"--- Page {page_num + 1} ---\n{text}\n\n"
            
            doc.close()
            
            if text_content.strip():
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text_content)
                return output_path
            else:
                raise Exception("No text content found in PDF")
                
        except Exception as e:
            raise Exception(f"Advanced PDF to text conversion failed: {str(e)}")
    
    async def _pdf_to_images_advanced(self, input_path, output_path, output_format):
        """Convert PDF to high-quality images (all pages)"""
        try:
            from pdf2image import convert_from_path
            import fitz
            
            # First, check if PDF has multiple pages
            doc = fitz.open(input_path)
            page_count = len(doc)
            doc.close()
            
            if page_count == 1:
                # Single page - convert directly
                images = convert_from_path(input_path, dpi=300, first_page=1, last_page=1)
                if images:
                    images[0].save(output_path, format=output_format.upper(), quality=95)
                    return output_path
            else:
                # Multiple pages - create a ZIP file with all pages
                import zipfile
                zip_path = output_path.rsplit('.', 1)[0] + '_all_pages.zip'
                
                images = convert_from_path(input_path, dpi=300)
                with zipfile.ZipFile(zip_path, 'w') as zipf:
                    for i, image in enumerate(images):
                        img_path = f"{output_path.rsplit('.', 1)[0]}_page_{i+1}.{output_format}"
                        image.save(img_path, format=output_format.upper(), quality=95)
                        zipf.write(img_path, f"page_{i+1}.{output_format}")
                        os.remove(img_path)  # Cleanup individual files
                
                return zip_path
                
            raise Exception("No pages found in PDF")
        except Exception as e:
            raise Exception(f"Advanced PDF to image conversion failed: {str(e)}")
    
    async def _pdf_to_docx_advanced(self, input_path, output_path):
        """Advanced PDF to DOCX conversion"""
        try:
            from pdf2docx import Converter
            
            cv = Converter(input_path)
            cv.convert(output_path, start=0, end=None)
            cv.close()
            
            # Verify conversion was successful
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                return output_path
            else:
                raise Exception("PDF to DOCX conversion produced empty file")
                
        except Exception as e:
            raise Exception(f"Advanced PDF to DOCX conversion failed: {str(e)}")
    
    async def _pdf_to_excel_advanced(self, input_path, output_path):
        """Advanced PDF to Excel conversion with table detection - ALL PAGES"""
        try:
            import fitz
            import pandas as pd
            import pdfplumber
            
            all_data = []
            
            # Try pdfplumber first for better table detection
            with pdfplumber.open(input_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables):
                        if table:  # Process all tables, even single row ones
                            # Clean the table data
                            cleaned_table = []
                            for row in table:
                                cleaned_row = [str(cell).strip() if cell is not None else "" for cell in row]
                                cleaned_table.append(cleaned_row)
                            
                            if cleaned_table:
                                # Use first row as header if we have multiple rows, otherwise create simple header
                                if len(cleaned_table) > 1:
                                    df = pd.DataFrame(cleaned_table[1:], columns=cleaned_table[0])
                                else:
                                    df = pd.DataFrame(cleaned_table, columns=[f"Column_{i+1}" for i in range(len(cleaned_table[0]))])
                                
                                all_data.append((f"Page_{page_num+1}_Table_{table_num+1}", df))
            
            # If no tables found, extract text from all pages
            if not all_data:
                doc = fitz.open(input_path)
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    text = page.get_text()
                    if text.strip():
                        # Create a simple DataFrame from text lines
                        lines = [line.strip() for line in text.split('\n') if line.strip()]
                        if lines:
                            df = pd.DataFrame(lines, columns=[f"Content_Page_{page_num+1}"])
                            all_data.append((f"Page_{page_num+1}", df))
                doc.close()
            
            if all_data:
                # Save all data to Excel with multiple sheets
                with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                    for sheet_name, df in all_data:
                        # Truncate sheet name if too long
                        sheet_name = sheet_name[:31]
                        df.to_excel(writer, sheet_name=sheet_name, index=False)
                
                logger.info(f"PDF to Excel conversion successful. Created {len(all_data)} sheets.")
                return output_path
            else:
                raise Exception("No extractable content found in PDF")
                
        except Exception as e:
            raise Exception(f"Advanced PDF to Excel conversion failed: {str(e)}")
    
    async def _text_to_pdf_advanced(self, input_path, output_path):
        """Advanced text to PDF conversion with professional formatting - FIXED HTML ISSUE"""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.units import inch
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            
            # Read text content
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                text_content = f.read()
            
            # Create professional PDF
            doc = SimpleDocTemplate(
                output_path,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            styles = getSampleStyleSheet()
            story = []
            
            # Custom style for better readability - SIMPLIFIED to avoid HTML parsing issues
            custom_style = ParagraphStyle(
                'Custom',
                parent=styles['Normal'],
                fontSize=12,
                leading=14,
                spaceAfter=12
            )
            
            # Split text into paragraphs and add to story - WITH HTML ESCAPING
            paragraphs = text_content.split('\n')
            for para in paragraphs:
                if para.strip():
                    # Escape HTML characters to prevent parsing issues
                    clean_para = html.escape(para.strip())
                    story.append(Paragraph(clean_para, custom_style))
                    story.append(Spacer(1, 12))
            
            if story:
                doc.build(story)
                return output_path
            else:
                raise Exception("No content to convert")
                
        except Exception as e:
            logger.error(f"Text to PDF conversion error: {e}")
            # Fallback to simple canvas method
            try:
                return await self._text_to_pdf_simple(input_path, output_path)
            except Exception as fallback_error:
                raise Exception(f"Advanced text to PDF conversion failed: {str(e)}")
    
    async def _text_to_pdf_simple(self, input_path, output_path):
        """Simple text to PDF conversion as fallback"""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            
            # Read text content
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                text_content = f.read()
            
            c = canvas.Canvas(output_path, pagesize=letter)
            width, height = letter
            
            # Set up text parameters
            y_position = height - 72  # Start from top with margin
            line_height = 14
            left_margin = 72
            
            # Split text into lines
            lines = []
            for paragraph in text_content.split('\n'):
                if paragraph.strip():
                    words = paragraph.split()
                    current_line = []
                    for word in words:
                        test_line = ' '.join(current_line + [word])
                        if c.stringWidth(test_line, "Helvetica", 10) < (width - 2 * left_margin):
                            current_line.append(word)
                        else:
                            if current_line:
                                lines.append(' '.join(current_line))
                            current_line = [word]
                    if current_line:
                        lines.append(' '.join(current_line))
                    lines.append('')  # Empty line between paragraphs
            
            # Draw text on canvas
            for line in lines:
                if y_position < 72:  # Bottom margin
                    c.showPage()
                    y_position = height - 72
                
                if line.strip():  # Only draw non-empty lines
                    c.drawString(left_margin, y_position, line)
                y_position -= line_height
            
            c.save()
            return output_path
            
        except Exception as e:
            raise Exception(f"Simple text to PDF conversion also failed: {str(e)}")
    
    async def _text_to_docx_advanced(self, input_path, output_path):
        """Advanced text to DOCX conversion"""
        try:
            from docx import Document
            from docx.shared import Pt
            
            with open(input_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            doc = Document()
            
            # Add professional styling
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)
            
            # Add content with proper paragraph formatting
            paragraphs = text_content.split('\n')
            for para in paragraphs:
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    p.paragraph_format.space_after = Pt(6)
            
            doc.save(output_path)
            return output_path
            
        except Exception as e:
            raise Exception(f"Advanced text to DOCX conversion failed: {str(e)}")
    
    async def _text_to_excel_advanced(self, input_path, output_path):
        """Convert text to Excel format"""
        try:
            import pandas as pd
            
            with open(input_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            # Split text into lines and create DataFrame
            lines = [line.strip() for line in text_content.split('\n') if line.strip()]
            df = pd.DataFrame(lines, columns=['Content'])
            
            # Save to Excel
            df.to_excel(output_path, index=False)
            return output_path
            
        except Exception as e:
            raise Exception(f"Text to Excel conversion failed: {str(e)}")
    
    async def _docx_to_pdf_advanced(self, input_path, output_path):
        """Advanced DOCX to PDF conversion using LibreOffice"""
        try:
            # Use LibreOffice for highest quality conversion
            cmd = [
                'libreoffice', '--headless', '--convert-to', 'pdf:writer_pdf_Export',
                '--outdir', os.path.dirname(output_path), input_path
            ]
            
            await self._run_command(cmd, timeout=120)
            
            # Find the converted file
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            possible_path = os.path.join(os.path.dirname(output_path), base_name + '.pdf')
            
            if os.path.exists(possible_path):
                if possible_path != output_path:
                    os.rename(possible_path, output_path)
                
                # Verify the PDF is valid
                if os.path.getsize(output_path) > 0:
                    return output_path
                else:
                    raise Exception("Conversion produced empty PDF")
            else:
                raise Exception("DOCX to PDF conversion failed - output not found")
                
        except Exception as e:
            raise Exception(f"Advanced DOCX to PDF conversion failed: {str(e)}")
    
    async def _docx_to_text_advanced(self, input_path, output_path):
        """Advanced DOCX to text conversion with formatting"""
        try:
            from docx import Document
            
            doc = Document(input_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Add table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            if text_content:
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write('\n'.join(text_content))
                return output_path
            else:
                raise Exception("No content found in DOCX file")
                
        except Exception as e:
            raise Exception(f"Advanced DOCX to text conversion failed: {str(e)}")
    
    async def _docx_to_excel_advanced(self, input_path, output_path):
        """Convert DOCX to Excel format"""
        try:
            from docx import Document
            import pandas as pd
            
            doc = Document(input_path)
            data = []
            
            # Extract paragraphs
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    data.append([f"Paragraph_{i+1}", paragraph.text.strip()])
            
            # Extract tables
            for table_num, table in enumerate(doc.tables):
                for row_num, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    if any(row_data):  # Only add non-empty rows
                        data.append([f"Table_{table_num+1}_Row_{row_num+1}"] + row_data)
            
            if data:
                # Create DataFrame with dynamic columns
                max_cols = max(len(row) for row in data)
                columns = ['Source'] + [f'Column_{i+1}' for i in range(max_cols - 1)]
                
                df = pd.DataFrame(data, columns=columns)
                df.to_excel(output_path, index=False)
                return output_path
            else:
                raise Exception("No content found in DOCX file")
                
        except Exception as e:
            raise Exception(f"DOCX to Excel conversion failed: {str(e)}")
    
    async def _excel_to_pdf_advanced(self, input_path, output_path):
        """Advanced Excel to PDF conversion"""
        try:
            # Use LibreOffice for best Excel to PDF conversion
            cmd = [
                'libreoffice', '--headless', '--convert-to', 'pdf:calc_pdf_Export',
                '--outdir', os.path.dirname(output_path), input_path
            ]
            
            await self._run_command(cmd, timeout=120)
            
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            possible_path = os.path.join(os.path.dirname(output_path), base_name + '.pdf')
            
            if os.path.exists(possible_path):
                if possible_path != output_path:
                    os.rename(possible_path, output_path)
                return output_path
            else:
                raise Exception("Excel to PDF conversion failed")
                
        except Exception as e:
            raise Exception(f"Advanced Excel to PDF conversion failed: {str(e)}")
    
    async def _excel_to_text_advanced(self, input_path, output_path):
        """Convert Excel to text format"""
        try:
            import pandas as pd
            
            # Read all sheets
            excel_file = pd.ExcelFile(input_path)
            text_content = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(input_path, sheet_name=sheet_name)
                text_content.append(f"--- Sheet: {sheet_name} ---")
                
                # Convert DataFrame to text
                for _, row in df.iterrows():
                    row_text = [str(cell) for cell in row if pd.notna(cell)]
                    if row_text:
                        text_content.append(" | ".join(row_text))
                
                text_content.append("")  # Empty line between sheets
            
            if text_content:
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write('\n'.join(text_content))
                return output_path
            else:
                raise Exception("No data found in Excel file")
                
        except Exception as e:
            raise Exception(f"Excel to text conversion failed: {str(e)}")
    
    async def _excel_to_docx_advanced(self, input_path, output_path):
        """Convert Excel to DOCX format"""
        try:
            import pandas as pd
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            
            # Read all sheets
            excel_file = pd.ExcelFile(input_path)
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(input_path, sheet_name=sheet_name)
                
                # Add sheet title
                doc.add_heading(f"Sheet: {sheet_name}", level=2)
                
                # Add table to document
                if not df.empty:
                    table = doc.add_table(rows=len(df)+1, cols=len(df.columns))
                    
                    # Add header row
                    for col_idx, column in enumerate(df.columns):
                        table.cell(0, col_idx).text = str(column)
                    
                    # Add data rows
                    for row_idx, (_, row) in enumerate(df.iterrows(), 1):
                        for col_idx, value in enumerate(row):
                            table.cell(row_idx, col_idx).text = str(value) if pd.notna(value) else ""
                
                doc.add_paragraph()  # Add space between sheets
            
            doc.save(output_path)
            return output_path
            
        except Exception as e:
            raise Exception(f"Excel to DOCX conversion failed: {str(e)}")
    
    async def _odt_to_pdf_advanced(self, input_path, output_path):
        """Advanced ODT to PDF conversion"""
        try:
            cmd = [
                'libreoffice', '--headless', '--convert-to', 'pdf:writer_pdf_Export',
                '--outdir', os.path.dirname(output_path), input_path
            ]
            
            await self._run_command(cmd, timeout=120)
            
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            possible_path = os.path.join(os.path.dirname(output_path), base_name + '.pdf')
            
            if os.path.exists(possible_path):
                if possible_path != output_path:
                    os.rename(possible_path, output_path)
                return output_path
            else:
                raise Exception("ODT to PDF conversion failed")
                
        except Exception as e:
            raise Exception(f"Advanced ODT to PDF conversion failed: {str(e)}")
    
    async def _convert_presentation(self, input_path, output_format, input_extension):
        """Professional presentation conversion"""
        try:
            output_path = input_path.rsplit('.', 1)[0] + f'.{output_format}'
            
            if output_format == 'pdf':
                if input_extension in ['pptx', 'ppt']:
                    return await self._ppt_to_pdf_advanced(input_path, output_path)
            
            raise Exception(f"Presentation conversion from {input_extension} to {output_format} not implemented")
            
        except Exception as e:
            raise Exception(f"Professional presentation conversion failed: {str(e)}")
    
    async def _ppt_to_pdf_advanced(self, input_path, output_path):
        """Advanced PowerPoint to PDF conversion"""
        try:
            cmd = [
                'libreoffice', '--headless', '--convert-to', 'pdf:impress_pdf_Export',
                '--outdir', os.path.dirname(output_path), input_path
            ]
            
            await self._run_command(cmd, timeout=180)
            
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            possible_path = os.path.join(os.path.dirname(output_path), base_name + '.pdf')
            
            if os.path.exists(possible_path):
                if possible_path != output_path:
                    os.rename(possible_path, output_path)
                return output_path
            else:
                raise Exception("PowerPoint to PDF conversion failed")
                
        except Exception as e:
            raise Exception(f"Advanced PowerPoint to PDF conversion failed: {str(e)}")
    
    async def _run_command(self, cmd, timeout=60):
        """Run system command with timeout"""
        try:
            process = await asyncio.create_subprocess_exec(
                *cmd,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE
            )
            
            stdout, stderr = await asyncio.wait_for(process.communicate(), timeout=timeout)
            
            if process.returncode != 0:
                error_msg = stderr.decode() if stderr else "Command failed"
                raise Exception(f"Command failed: {error_msg}")
            
            return stdout.decode() if stdout else ""
            
        except asyncio.TimeoutError:
            raise Exception(f"Command timeout after {timeout} seconds")
        except Exception as e:
            raise Exception(f"Command execution failed: {str(e)}")

# Global universal converter instance
universal_converter = UniversalConverter()